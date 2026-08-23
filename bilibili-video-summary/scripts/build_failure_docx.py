"""Build a Word report when a video could not be processed successfully."""
import argparse
import json
from pathlib import Path


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--failure", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--metadata")
    args = parser.parse_args()

    try:
        from docx import Document
    except ImportError as exc:
        raise SystemExit("FAIL: python-docx is not installed") from exc

    from build_notes_docx import configure_document, set_east_asia_font

    failure_path = Path(args.failure)
    failure = json.loads(failure_path.read_text(encoding="utf-8-sig"))
    if not isinstance(failure, dict):
        raise SystemExit("FAIL: failure record must be a JSON object")
    required = ("stage", "reason", "attempted_methods", "suggested_action")
    missing = [key for key in required if not failure.get(key)]
    if missing:
        raise SystemExit(f"FAIL: failure record is missing fields: {missing}")

    metadata = {}
    if args.metadata:
        metadata = json.loads(Path(args.metadata).read_text(encoding="utf-8-sig"))
        if not isinstance(metadata, dict):
            raise SystemExit("FAIL: metadata must be a JSON object")

    document = Document()
    configure_document(document)
    title = metadata.get("title") or metadata.get("bvid") or "B站视频处理失败记录"
    document.core_properties.title = f"{title} - 处理失败记录"
    document.add_heading("B站视频处理失败记录", 0)
    for label, value in (
        ("标题", metadata.get("title")),
        ("BV号", metadata.get("bvid")),
        ("来源", metadata.get("source_url")),
        ("失败阶段", failure.get("stage")),
    ):
        if value:
            paragraph = document.add_paragraph()
            label_run = paragraph.add_run(f"{label}：")
            label_run.bold = True
            set_east_asia_font(label_run, "Microsoft YaHei")
            value_run = paragraph.add_run(str(value))
            set_east_asia_font(value_run, "Microsoft YaHei")

    document.add_heading("失败原因", level=1)
    document.add_paragraph(str(failure["reason"]))
    document.add_heading("已尝试方法", level=1)
    for method in failure["attempted_methods"]:
        document.add_paragraph(str(method), style="List Bullet")
    document.add_heading("建议动作", level=1)
    document.add_paragraph(str(failure["suggested_action"]))
    document.add_paragraph("本文件记录处理失败，不代表视频内容已完成总结。")

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    visual_qa = output.with_name("visual-qa.json")
    visual_qa.write_text(
        json.dumps(
            {
                "status": "unverified",
                "reason": "failure report has not been rendered and visually inspected",
                "renderer": None,
                "page_count": None,
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )
    print(
        json.dumps(
            {"status": "recorded_failure", "output": output.name},
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
