"""Build a consistent Chinese Word note from analysis and cleaned subtitles."""
import argparse
import json
import re
from pathlib import Path


def read_text(path: Path) -> str:
    text = path.read_text(encoding="utf-8-sig").strip()
    if not text:
        raise SystemExit(f"FAIL: empty input file: {path}")
    return text


def load_uncertain_terms(path: Path | None) -> list[str]:
    if path is None:
        return []
    if path.suffix.lower() == ".json":
        value = json.loads(path.read_text(encoding="utf-8-sig"))
        if not isinstance(value, list) or not all(isinstance(item, str) for item in value):
            raise SystemExit("FAIL: uncertain terms JSON must be a string array")
        terms = value
    else:
        terms = path.read_text(encoding="utf-8-sig").splitlines()
    return sorted({term.strip() for term in terms if term.strip()}, key=len, reverse=True)


def set_east_asia_font(run, font_name: str) -> None:
    from docx.oxml.ns import qn

    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)


def add_marked_text(paragraph, text: str, uncertain_terms: list[str]) -> None:
    from docx.shared import RGBColor

    if not uncertain_terms:
        run = paragraph.add_run(text)
        set_east_asia_font(run, "Microsoft YaHei")
        return
    pattern = re.compile("(" + "|".join(re.escape(term) for term in uncertain_terms) + ")")
    uncertain = set(uncertain_terms)
    for piece in pattern.split(text):
        if not piece:
            continue
        run = paragraph.add_run(piece)
        set_east_asia_font(run, "Microsoft YaHei")
        if piece in uncertain:
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            run.font.underline = True


def configure_document(document) -> None:
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(4)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_metadata_table(document, metadata: dict) -> None:
    from docx.shared import Cm

    rows = [
        ("BV号", metadata.get("bvid")),
        ("标题", metadata.get("title")),
        ("UP主", metadata.get("owner")),
        ("时长", metadata.get("duration_seconds")),
        ("字幕来源", metadata.get("subtitle_source")),
        ("来源链接", metadata.get("source_url")),
    ]
    rows = [(label, value) for label, value in rows if value not in (None, "")]
    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for index, (label, value) in enumerate(rows):
        table.cell(index, 0).width = Cm(3.2)
        table.cell(index, 1).width = Cm(12.0)
        table.cell(index, 0).text = str(label)
        table.cell(index, 1).text = str(value)
        for run in table.cell(index, 0).paragraphs[0].runs:
            run.bold = True
            set_east_asia_font(run, "Microsoft YaHei")
        for run in table.cell(index, 1).paragraphs[0].runs:
            set_east_asia_font(run, "Microsoft YaHei")


def add_markdown(document, content: str, uncertain_terms: list[str]) -> None:
    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        if not line:
            document.add_paragraph()
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            paragraph = document.add_heading(level=len(heading.group(1)))
            add_marked_text(paragraph, heading.group(2), uncertain_terms)
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            add_marked_text(paragraph, bullet.group(1), uncertain_terms)
            continue
        numbered = re.match(r"^\d+[.]\s+(.+)$", line)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_marked_text(paragraph, numbered.group(1), uncertain_terms)
            continue
        paragraph = document.add_paragraph()
        add_marked_text(paragraph, line, uncertain_terms)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--metadata", required=True)
    parser.add_argument("--analysis", required=True)
    parser.add_argument("--subtitle", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--uncertain-terms")
    args = parser.parse_args()

    try:
        from docx import Document
    except ImportError as exc:
        raise SystemExit("FAIL: python-docx is not installed") from exc

    metadata_path = Path(args.metadata)
    metadata = json.loads(metadata_path.read_text(encoding="utf-8-sig"))
    if not isinstance(metadata, dict):
        raise SystemExit("FAIL: metadata must be a JSON object")
    for required in ("bvid", "title"):
        if not str(metadata.get(required, "")).strip():
            raise SystemExit(f"FAIL: metadata is missing {required}")

    analysis = read_text(Path(args.analysis))
    subtitle = read_text(Path(args.subtitle))
    uncertain_terms = load_uncertain_terms(
        Path(args.uncertain_terms) if args.uncertain_terms else None
    )

    document = Document()
    configure_document(document)
    document.core_properties.title = str(metadata["title"])
    document.core_properties.subject = "Bilibili video summary"
    document.add_heading(str(metadata["title"]), 0)
    add_metadata_table(document, metadata)
    document.add_paragraph()
    add_markdown(document, analysis, uncertain_terms)
    document.add_page_break()
    document.add_heading("整理阅读版字幕", level=1)
    add_markdown(document, subtitle, uncertain_terms)

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    visual_qa = output.with_name("visual-qa.json")
    visual_qa.write_text(
        json.dumps(
            {
                "status": "unverified",
                "reason": "document has not been rendered and visually inspected",
                "renderer": None,
                "page_count": None,
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )
    metadata["status"] = "completed"
    metadata["notes_document"] = "analysis/notes.docx"
    metadata["visual_qa"] = "analysis/visual-qa.json"
    transcription = metadata_path.parent / "subtitles" / "transcription.json"
    if transcription.exists():
        metadata["subtitle_source"] = "local_whisper"
        metadata["transcription"] = "subtitles/transcription.json"
    metadata_path.write_text(
        json.dumps(metadata, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    print(
        json.dumps(
            {
                "status": "success",
                "output": output.name,
                "bytes": output.stat().st_size,
                "uncertain_term_count": len(uncertain_terms),
                "visual_qa": "unverified",
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
